from django.views.generic import ListView
from django.http import Http404
from .models import User, Attendance
from .forms import UserSearchForm
from django.http import HttpResponse
import urllib
import csv
import os
from pathlib import Path
from django.db.models import Q
import datetime
import locale
from docx import Document

class AttendanceListView(ListView):
    model = Attendance
    template_name = 'attendance_list.html'
    paginate_by = 20
    
    def post(self, request, *args, **kwargs):
        return self.get(request, *args, **kwargs)

    def get_queryset(self):
        q_name = self.request.POST.get('name')
        q_date = self.request.POST.get('date')
        dt = None

        if q_date:
            date = q_date.split('-')
            try:
                dt = datetime.datetime(int(date[0]), int(date[1]), int(date[2]))
            except Exception as e:
                pass

        return filter_attendance(q_name, dt)

def filter_attendance(name, dt):
    if dt:
        object_list = Attendance.objects.filter(
            Q(user__name__icontains=name) & 
            Q(accepted_at__gte=dt) &
            Q(accepted_at__lte=dt+datetime.timedelta(days = 1))
        )
    elif name:
        object_list = Attendance.objects.filter(
            Q(user__name__icontains=name)
        )
    else:
        object_list = Attendance.objects.all()

    return object_list

def PaperExport(request):
    # Null check
    if not request.POST.get('date'):
        raise Http404("Date does not exist")

    # 日付計算
    current_date_str = request.POST.get('date')
    current_dt = datetime.datetime.strptime(current_date_str, '%Y-%m-%d')
    # 記録開始から何周目なのか
    week_count = int((current_dt.date() - datetime.datetime(2020,6,1).date()).days / 7 + 1)
    # 対象週の月曜日
    monday = current_dt - datetime.timedelta(days=current_dt.weekday())

    # Generate docx
    filepath = f'申請書_week{week_count}.docx'
    document = Document(Path(os.environ['DOCX_PATH'], 'template.docx'))

    # Collect attendances
    for i in range(7):
        target_dt = monday + datetime.timedelta(days=i)
        attendances = filter_attendance('', target_dt)
        member_set = {}
        for att in attendances:
            entrance_time_checker(member_set, att)

        paper_1day_writer(document, target_dt, member_set)

    document.save(Path(os.environ['DOCX_PATH'], filepath))

    with open(Path(os.environ['DOCX_PATH'], filepath), 'rb') as fp:
        data = fp.read()

    # Generate response file
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = urllib.parse.quote(filepath)
    print(filename)
    response['Content-Disposition'] = f"attachment; filename='{filename}'; filename*=UTF-8''{filename}"
    response.write(data)

    return response

def paper_1day_writer(document, dt, member_set):
    document.tables[dt.weekday()].rows[1].cells[0].text = str(dt.month) + '月' + str(dt.day) + '日'
    document.tables[dt.weekday()].rows[1].cells[0].text += f'（{weekday_fig(dt.weekday())}）'
    for i, member in enumerate(member_set.values()):
        document.tables[dt.weekday()].rows[1+i].cells[1].text = time_fig(member[0], member[1])
        document.tables[dt.weekday()].rows[1+i].cells[2].text = member[2]
        document.tables[dt.weekday()].rows[1+i].cells[3].text = member[3]

def entrance_time_checker(member_set, attendance):
    # 記録されている時間
    rec_time = attendance.accepted_at.astimezone(
        datetime.timezone(datetime.timedelta(hours=+9))
        ).time()
    # 時間計算
    if attendance.is_in:
        is_am = rec_time < datetime.time(12, 30)
    else:
        is_am = rec_time < datetime.time(13, 15)

    # member_setに記録されていない学生の場合
    if not attendance.user.number in member_set.keys():
        # Word内の表に対応した順：午前,午後,名前,学生証番号
        member_set[attendance.user.number] = [
            is_am,
            not is_am,
            attendance.user.name,
            attendance.user.number
        ]
    # 記録済だった場合
    else:
        if member_set[attendance.user.number][0] and not is_am or \
        not member_set[attendance.user.number][0] and is_am:
            member_set[attendance.user.number][0] = True
            member_set[attendance.user.number][1] = True

def time_fig(am, pm):
    if am and pm:
        return '終日'
    elif am and not pm:
        return '午前'
    else:
        return '午後'

def weekday_fig(weekday):
    if weekday == 0:
        return '月'
    elif weekday == 1:
        return '火'
    elif weekday == 2:
        return '水'
    elif weekday == 3:
        return '木'
    elif weekday == 4:
        return '金'
    elif weekday == 5:
        return '土'
    elif weekday == 6:
        return '日'
    else:
        return 'X'
